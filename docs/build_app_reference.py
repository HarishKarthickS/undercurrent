from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name("Undercurrent_App_Reference.docx")

NAVY = "18314F"
TEAL = "287D7D"
GOLD = "B7791F"
PALE = "EEF5F4"
GRAY = "5B6573"
BLACK = "1E2933"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
for style_name, size, color, bold, before, after in [
    ("Title", 28, NAVY, True, 0, 8), ("Subtitle", 12, GRAY, False, 0, 14),
    ("Heading 1", 18, NAVY, True, 18, 7), ("Heading 2", 13, TEAL, True, 12, 5),
    ("Heading 3", 11, NAVY, True, 8, 3), ("Normal", 10.2, BLACK, False, 0, 6),
]:
    st = styles[style_name]
    st.font.name = "Aptos"; st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    st.font.size = Pt(size); st.font.bold = bold; st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.15

for s in ("List Bullet", "List Number"):
    styles[s].font.name = "Aptos"; styles[s].font.size = Pt(10.2)
    styles[s].paragraph_format.space_after = Pt(3)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None: node = OxmlElement(f"w:{side}"); tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h; shade(c, NAVY); set_cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in c.paragraphs[0].runs: run.font.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(9)
        if widths: c.width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value); set_cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: cells[i].width = Inches(widths[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
                for run in p.runs: run.font.size = Pt(8.7)
    doc.add_paragraph()
    return t

def p(text="", style=None, bold_lead=None):
    para = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r = para.add_run(bold_lead); r.bold = True; para.add_run(text[len(bold_lead):])
    else: para.add_run(text)
    return para

def bullets(items):
    for x in items: p(x, "List Bullet")

def page_break(): doc.add_page_break()

def callout(label, text):
    t = doc.add_table(rows=1, cols=1); t.autofit = False; c=t.cell(0,0); shade(c, PALE); set_cell_margins(c, 160, 180, 160, 180)
    q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(3); r=q.add_run(label.upper()); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(TEAL)
    q=c.add_paragraph(text); q.paragraph_format.space_after=Pt(0)
    doc.add_paragraph()

def heading(text, level=1): doc.add_heading(text, level=level)

# Cover
p("UNDERCURRENT", "Title")
p("Complete Application Reference", "Subtitle")
p("Product, experience, architecture, technology, data, parent and child flows, and page-by-page guide", "Subtitle")
callout("Document purpose", "A single implementation-aware reference for the current Undercurrent codebase. It explains what the app does, how the parts collaborate, and what each user-facing route and dashboard area is responsible for.")
p(f"Prepared from the repository source and internal documentation on {date.today().isoformat()}.")
heading("At a glance", 1)
table(["Area", "What it is"], [
    ["Product", "A parent-visible learning companion where a child teaches Pip about an idea from their day."],
    ["Child experience", "A parent-approved device, PIN-gated student trail with morning rituals and evening learning moments."],
    ["Parent experience", "A private control center for routines, device approval, family preferences, summaries, signals, and conversation support."],
    ["Core promise", "Process over performance: no report card, grades, mastery claims, streak pressure, or child-facing parent dashboard."],
    ["Deployment boundary", "Closed demo only; public enrollment is intentionally rejected."],
], [1.5, 5.0])
heading("Contents", 1)
for x in ["1. Product and core concepts", "2. Technology and architecture", "3. Data, security, and safety", "4. Parent flow", "5. Child flow", "6. Page-by-page guide", "7. API and operations reference", "8. Current constraints and implementation notes"]: p(x, "List Number")

page_break()
heading("1. Product and core concepts")
p("Undercurrent is a family learning companion designed around a short, low-pressure exchange: a child teaches the curious character Pip about something from their day. Pip responds with one age-appropriate follow-up that supports explanation, reflection, and curiosity rather than judging correctness.")
heading("Primary users", 2)
table(["User", "Needs", "What the app provides"], [
    ["Parent / guardian", "Set up the household, decide how the experience fits family life, and stay lightly informed.", "An authenticated parent center with child profiles, settings, safe summaries, device controls, and optional conversation support."],
    ["Student / child", "A simple private place to start a ritual, share an idea, and return later without pressure.", "A device-approved, PIN-unlocked trail with ritual cards, guided quests or open chat, voice affordances, and gentle closure."],
    ["Operator", "Control demo enrollment and deploy the system safely.", "Invitation-bound parent signup, service configuration, migrations, Docker topology, and operational documentation."],
], [1.35, 2.55, 2.6])
heading("Core learning model", 2)
bullets([
    "Morning Ripple is a lightweight, four-stage daily ritual. It is curated by default, may collect mood/energy/path/intention/collectible information, and does not require the AI service.",
    "Evening Discovery is the ordinary teaching moment. The child gives an explanation, the service safety-screens it, then (when configured) uses structured AI assessment and a constrained follow-up response.",
    "Topics, spaced-review dates, heuristic understanding/confidence signals, effort moments, and cosmetic wins help shape later invitations. These are descriptive signals, not grades or diagnoses.",
    "Parent visibility is deliberately summary-led. Safety-trigger content is excluded from recaps, transcript views, digests, and AI advisor context.",
])
callout("Design principle", "The app is built for small, optional, repeatable moments. Routine settings change invitations and boundaries; they do not turn exploration into a performance target.")
heading("Session rules", 2)
table(["Rule", "Morning", "Evening"], [
    ["Opening", "A short arrival/feeling prompt.", "An invitation to teach Pip one small thing."],
    ["AI requirement", "No. Curated content stays available.", "Yes for ordinary AI-backed turns; otherwise the API fails closed with AI_UNAVAILABLE."],
    ["Content storage", "Completed ritual entry and metadata.", "Ordinary child and companion turns encrypted at rest; safety text is not stored."],
    ["Completion", "Ritual completion returns a keepsake/entry.", "Ends after a safety event, the configured cap, or explicit child exit."],
], [1.5, 2.45, 2.55])

page_break()
heading("2. Technology and architecture")
p("The repository is an npm workspace with a feature-oriented React frontend and a module-oriented Fastify backend. The browser uses a same-origin /api boundary; the proxy removes the /api prefix before requests reach Fastify.")
heading("Technology stack", 2)
table(["Layer", "Technology", "Role"], [
    ["Web", "React 19, Vite 8, plain CSS", "Single-page application, route selection, feature components, and API adapters."],
    ["API", "Node.js, Fastify 5", "HTTP routes, validation, authentication, domain-service composition, and error handling."],
    ["Persistence", "PostgreSQL 17, Drizzle ORM / drizzle-kit", "System of record, schema definitions, migrations, and repositories."],
    ["Cache / protection", "Redis 7, Fastify rate-limit, under-pressure", "Rate-limit store, readiness dependencies, and overload protection."],
    ["AI", "OpenAI JS SDK, Responses API-compatible provider", "Two structured private assessments and one constrained companion reply for ordinary evening turns."],
    ["Security", "Argon2, AES-256-GCM, SHA-256 token hashes, Helmet, CSRF/cookie tooling", "Password protection, encrypted turns, secure parent session handling, HTTP hardening."],
    ["Notifications", "Nodemailer / SMTP", "Parent safety alerts; JSON transport is used when SMTP is absent."],
    ["Delivery", "Docker Compose, Nginx, GitHub workflows", "Local/production-like composition, static-web hosting, reverse proxy, and CI support."],
], [1.2, 2.35, 3.0])
heading("Frontend organization", 2)
table(["Location", "Responsibility"], [
    ["apps/web/src/app", "App shell, browser-history routing hook, main route switch, and shared application controller."],
    ["features/landing", "Public product landing page and conversion entry points."],
    ["features/identity", "Parent authentication plus student device invitation, picker, PIN setup, and unlock views."],
    ["features/onboarding", "Child profile directory and three-step creation flow."],
    ["features/child-session", "Student home, Morning Ripple, guided/open session UI, voice helpers, and session APIs."],
    ["features/parent-dashboard", "Parent tabs, summary/detail pages, routines, device controls, advisor, and parent APIs."],
    ["shared", "HTTP client, async state helper, API conventions, icons, mascot, and reusable visual components."],
], [2.5, 4.05])
heading("Backend organization", 2)
table(["Location", "Responsibility"], [
    ["bootstrap/createServer.js", "Composition root: infrastructure clients, repositories, services, middleware, public routes, and protected scope."],
    ["modules/identity", "Invitation-bound account lifecycle, parent sessions, demo terms, guardian invitations."],
    ["modules/students", "Household-scoped student listing and creation with consent enforcement."],
    ["modules/student-access", "Approved-device invitation, PIN, device tokens, unlock tokens, revoke/reset controls."],
    ["modules/sessions", "Morning/evening lifecycle, idempotency, safety-first turn processing, history, and settings."],
    ["modules/learning and modules/ai", "Assessment reconciliation, scaffolding, review scheduling, rewards, and constrained model calls."],
    ["modules/parent-dashboard and privacy", "Parent summaries, advisor, experience preferences, export, withdrawal, and deletion."],
    ["platform", "Database schemas/repositories, encryption, HTTP middleware, error model, Redis, and mailer."],
], [2.5, 4.05])
heading("Request path", 2)
p("Browser -> Vite development proxy or Nginx -> Fastify route -> authentication/validation/middleware -> domain service -> repository and infrastructure adapters -> JSON response. Protected parent routes attach a parent session; student routes derive a student actor from device and unlock cookies.")

page_break()
heading("3. Data, security, and safety")
heading("Key data relationships", 2)
table(["Entity group", "Important records", "Why it exists"], [
    ["Household identity", "households, parent_accounts, parent_sessions", "Defines the tenant boundary, parent credentials, and 12-hour parent sessions."],
    ["Student access", "students, student access/device records, consent_records", "Holds profile information, approved-device controls, PIN/unlock lifecycle, and learning-companion consent."],
    ["Learning sessions", "sessions, turn_requests, session_turns", "Tracks session lifecycle, deduplicates retries by session/idempotency key/input hash, and stores encrypted ordinary turns."],
    ["Learning signals", "topics, scores, wins", "Maintains learned-topic metadata, heuristic assessments, spaced review, and effort/collectible outcomes."],
    ["Safety and operations", "safety_events, outbox_events, product_analytics_events", "Records safety metadata without triggering text, future durable events, and allowlisted opt-in setup analytics."],
], [1.5, 2.65, 2.4])
heading("Security controls", 2)
bullets([
    "Parent passwords are hashed with Argon2id. Parent session tokens are stored as SHA-256 hashes and expire after 12 hours.",
    "Ordinary evening messages and replies are encrypted using AES-256-GCM with a unique 12-byte IV, authentication tag, and stored key version for every turn.",
    "The API applies Helmet headers, JSON-only body enforcement, body-size limits, request IDs, Redis-backed IP rate limits, and pressure checks.",
    "Every protected service action checks household ownership and the student's granted learning-companion consent before proceeding.",
    "The child device uses parent-approved device tokens plus a 4-8 digit PIN and a 12-hour unlock cookie; the web app also preserves a local 12-hour unlock convenience state.",
])
heading("Safety-first evening turn flow", 2)
table(["Step", "System behavior", "Privacy consequence"], [
    ["1. Receive", "Session service validates actor, session, limits, and idempotency key.", "Duplicate completed work can be replayed only for the same hashed input."],
    ["2. Screen", "Deterministic regular expressions check narrowly for immediate danger, unsafe-at-home, and medical-emergency signals.", "This check runs before encryption and before any AI provider call."],
    ["3. Escalate if matched", "Create safety-event metadata, end the session, send a parent alert, and return a trusted-grown-up redirect.", "The triggering words are never inserted into session_turns or surfaced in parent summaries."],
    ["4. Process clear turn", "Require usable AI and encryption services; encrypt child input; run two structured assessments; reconcile learning signal; generate constrained reply.", "If the required services are unavailable, return AI_UNAVAILABLE rather than fabricating a personalized answer."],
    ["5. Persist and respond", "Encrypt Pip's reply, persist topic/score/review/win metadata, then respond to the child UI.", "Parents receive summary data; transcript access is separately controlled in the family settings."],
], [0.85, 3.2, 2.5])
callout("Important limitation", "This is a closed-demo implementation, not a public child service or a verified-parental-consent solution. Production use requires legal, privacy, safeguarding, monitoring, delivery, and key-rotation work described in the repository's security and operations guides.")

page_break()
heading("4. Parent flow - detailed")
p("The parent flow owns family setup, policy choices, visibility, and device handoff. It is intentionally separate from the student experience: parents do not enter a child's learning screen through the parent center.")
heading("Parent lifecycle", 2)
table(["Stage", "User action", "System behavior", "Outcome"], [
    ["Invitation and account", "Open an operator-issued sign-up link, create credentials, then verify email.", "Identity service validates invitation and creates account/session; parent session is delivered via secure cookie.", "Authenticated guardian can enter the family control center."],
    ["Child profile", "Add a child through the three-step setup modal.", "Terms acknowledgement is recorded; student plus granted learning-companion consent are created; ritual settings are saved.", "New profile appears in the children directory and parent tabs."],
    ["Device handoff", "Send an approved-device invitation from Family.", "A one-time invitation is consumed on the device, which receives a device token; child sets a PIN.", "Student can unlock only their approved trail on that device."],
    ["Family rhythm", "Set hours, session limit, voice/activity availability, interaction band, and preferred mode.", "Ritual/session settings affect later child invitations and presentation.", "A bounded, tailored starting experience."],
    ["Ongoing visibility", "Review Today, rituals, conversations, learning, insights, and optionally Ask Pip.", "Dashboard aggregates metadata, summaries, and configured access; safety needs take precedence.", "Parent gets context without treating the child as a score."],
], [1.1, 1.9, 2.65, 0.9])
heading("Parent authorization flow", 2)
p("On app startup, useAppController calls the parent-session endpoint. If a cookie-backed session returns a parent, the controller stores a parent token/object for protected API calls and loads household profiles. The app route switch prevents access to /parent areas until this readiness check completes. Missing authentication redirects to the parent login view.")
heading("What a parent can control", 2)
bullets([
    "Profiles: create a child profile, move between household children, and see setup nudge progress.",
    "Rhythm: morning/evening start hours, daily session limit, voice responses, tiny activities, interaction band, and preferred experience style.",
    "Morning studio: optional AI prompts, available activity paths, sensitivity tone, and completed ritual-entry history.",
    "Devices: send approved-device link, review device list/last activity, revoke one device, reset PIN, or sign out all devices.",
    "Visibility and privacy: daily/weekly recaps, aggregate setup analytics consent, transcript access, and transcript-aware parent advisor consent.",
    "Safety: acknowledge a dashboard safety follow-up. The parent sees category/action context, not the child's triggering text.",
])

page_break()
heading("5. Child flow - detailed")
p("The child flow begins only after a parent has approved the device. It keeps the student in a persistent, child-facing shell rather than routing a returning student back to the public marketing page.")
heading("Student access and session lifecycle", 2)
table(["Stage", "Child experience", "Behind the scenes"], [
    ["1. Device invitation", "Opening the one-time parent link shows a short setup state.", "The invitation is consumed; a device token is added to an HTTP-only cookie."],
    ["2. Profile selection", "The device lists only approved student profiles. The child chooses their own trail.", "The API resolves profiles from device tokens; no parent credentials are used."],
    ["3. PIN setup/unlock", "First use chooses a 4-8 digit PIN; later use enters the PIN.", "Server grants an unlock token for 12 hours; client retains a corresponding local convenience state."],
    ["4. Student home", "Pip presents the next gentle invitation based on time, grade/band, settings, and available modes.", "Home/session APIs calculate daily world/ritual context and enforce parent-configured boundaries."],
    ["5. Morning Ripple", "Child progresses through arrival, feeling/energy, a selected path, and a launch/intention card.", "Completion saves a ritual entry with selected metadata and cosmetic collectible; normal flow does not need AI."],
    ["6. Evening session", "Child chooses a guided quest or chat and can type, tap, or use voice when allowed.", "Every submitted turn receives an idempotency key; service applies safety, AI/encryption, learning persistence, and terminal rules."],
    ["7. Closing", "Child can save for later; terminal/safety states show a calm return path.", "Session end reason is recorded. Safety terminal flow has fixed trusted-grown-up copy and parent alert behavior."],
], [1.2, 2.65, 2.65])
heading("Voice and interaction behavior", 2)
bullets([
    "The client uses browser recognition and synthesis capabilities through the useVoice hook. Voice is disabled when parent settings turn it off or browser support is absent.",
    "The session UI exposes state such as listening, thinking, speaking, fallback, completed, and safety-ended. Controls are disabled while work is in progress.",
    "In guided mode, a quest can supply title, step label, and progress copy. In chat mode, TalkToPip presents a message-focused interface with the same session lifecycle.",
    "Morning and evening are intentionally different: morning is a short ritual; evening supports the safety-screened teaching exchange.",
])
callout("Child-centered boundary", "Pip is prompted to use age-appropriate, process-focused language and to avoid grades, correctness declarations, emotional dependency, friendship claims, and therapeutic authority.")

page_break()
heading("6. Page-by-page guide")
heading("Public and account pages", 2)
table(["Route", "Page", "What it shows and does"], [
    ["/", "Landing page", "Public field-guide style introduction to the product. Calls to action route to parent setup/sign-in. If the device has an active student unlock, the same root route opens the student home instead."],
    ["/parent/signup?invite=...", "Parent sign-up", "Closed-demo invitation-bound account creation. Collects name, email, and 12+ character password; then presents email-verification confirmation."],
    ["/parent/login", "Parent login", "Email/password sign-in. A successful session routes to the children directory."],
    ["/parent/forgot-password", "Password reset request", "Collects email and displays a neutral reset-delivery status."],
    ["/parent/verify", "Email verification action", "Consumes parent ID/token query parameters and opens the authenticated control center on success."],
    ["/parent/reset-password", "Password reset completion", "Accepts reset token and a new 12+ character password."],
    ["/parent", "Parent entry", "Minimal authenticated entry page that points to Manage children."],
    ["/parent/children", "Children directory and onboarding", "Shows existing approved profiles; Add a child opens a three-step modal for identity, rhythm, and starting interaction band."],
], [1.85, 1.55, 3.1])
heading("Student-device pages", 2)
table(["Route", "Page", "What it shows and does"], [
    ["/student/invite/:token", "Invitation consumer", "Consumes a one-time parent-approved device link, registers device access, and returns to the student picker."],
    ["/student", "Student picker", "Shows only profiles attached to the approved device. If none exist, explains that a parent must send an invitation."],
    ["/student/:id/setup", "PIN setup", "Creates a numeric 4-8 digit PIN for a profile on the approved device."],
    ["/student/:id/unlock", "PIN unlock", "Checks the student PIN and opens the student trail."],
    ["/student/:id", "Student home", "Renders StudentRitual when no session is active; renders the current child session once started."],
], [1.85, 1.55, 3.1])
heading("Student home and session screens", 2)
table(["Screen", "Detail"], [
    ["StudentRitual home", "Displays Pip's daily invitation and available modes. Starting a session calls the student session-start API with session type, mode, and optional quest ID."],
    ["MorningRipple", "Four-stage ritual flow: arrival, feeling/energy selection, rotating path/activity, and launch/intention/collectible completion. Completion is saved through the student-only API."],
    ["Guided ChildSession", "Hero, Pip response bubble, optional text-to-speech replay, chips/text/voice input, progress/status copy, and Save for later. Handles thinking/listening/speaking/terminal states."],
    ["TalkToPip", "Open-chat presentation of messages and text/voice submission controls for chat-mode sessions."],
    ["TerminalPanel", "Calm completion or safety-ended boundary view with a return action. It does not expose sensitive safety detail."],
], [1.8, 4.7])

page_break()
heading("Parent Center tabs")
table(["Route suffix", "Tab", "Detailed responsibilities"], [
    ["today", "Today", "Shows next gentle step or safety check-in, daily/weekly moment counts, topic/effort counts, routine detail, setup nudge, optional family conversation starter, recent moments, and a family glance switcher."],
    ["rituals", "Rituals", "Morning Ripple studio: reads completed entries and lets the parent configure optional AI prompts, available paths, and sensitivity tone."],
    ["conversations", "Conversations", "Parent-only ordinary-session recap/archive area. It can list conversations and open eligible sessions. Safety-trigger content is excluded."],
    ["learning", "Learning", "Topic-focused detail view with learning signals, review context, and child-safe interpretation framing rather than evaluation."],
    ["insights", "Insights", "Topic trail, review dates, gentle trend labels, effort moments, and suggested family questions. It explicitly frames data as clues, not scores."],
    ["advisor", "Ask Pip", "Parent-only question interface. Provides starter questions, advisor history, custom question entry, clear-history action, and summary-only/no-child-facing-change guardrails."],
    ["plan", "Plan", "Controls morning/evening hours, daily boundary, voice/tiny activities, interaction band override, and preferred session style. Saves settings for the next child visit."],
    ["family", "Family", "Approved-device handoff, device management, PIN reset, sign-out everywhere, daily/weekly digest preferences, opt-in aggregate analytics, transcript access, and advisor consent."],
], [1.25, 1.25, 4.0])
p("Parent URL pattern: /parent/children/:studentId/:section. The router aliases customize, inbox, and privacy to the Family tab, while recognized tabs drive the ParentDashboard component. A left rail and mobile bottom navigation expose the same tab set.")
heading("Parent dashboard data contract", 2)
p("The dashboard hook fetches a selected student's summary once the parent identity and student profile are available. The UI expects student identity, home/next action, weekly counts, topic list, effort moments, safety context, conversation starters, recent moments, household children, setup progress, controls, daily brief, and settings-related data. A Refresh family signals control reloads the dashboard after actions such as safety acknowledgement or settings changes.")

page_break()
heading("7. API and operations reference")
heading("API groups", 2)
table(["Group", "Representative endpoints", "Authorization"], [
    ["Health", "GET /health, GET /ready", "Public."],
    ["Identity", "POST /auth/signup, /auth/login, /auth/verify-email, /auth/logout; GET /auth/session; terms, session revocation, guardian invitations", "Public for account lifecycle; parent session required for terms/session administration."],
    ["Students", "GET/POST /students", "Parent session and household ownership."],
    ["Student access", "Invitation consume, device profiles, PIN set/unlock; parent invitation/access/revoke/reset endpoints", "Device cookies for child actions; parent session for management."],
    ["Sessions", "Start, turn, end, curiosity trail, ritual settings, reflections, conversations, Morning Ripple history", "Parent actor or separately registered student actor depending on route family."],
    ["Dashboard", "Dashboard get/update, household experience, safety acknowledge, topic detail, advisor history/turn", "Parent session and student household ownership."],
    ["Privacy", "Student export, withdraw, delete", "Parent session and student household ownership."],
], [1.35, 3.85, 1.3])
heading("Error and retry behavior", 2)
bullets([
    "Expected API errors use an error object with code, message, and requestId. Unexpected errors are logged server-side and respond with a generic internal-error message.",
    "The shared web HTTP client uses a 30-second abort timeout. AI-backed server work can outlive that browser timeout, so a retry of the same input must retain the same idempotency key.",
    "A completed idempotent session turn may replay only when session ID, key, and input hash agree; mismatched or in-progress duplicate requests return conflict behavior.",
])
heading("Local development and deployment", 2)
table(["Task", "Command / component", "Notes"], [
    ["Start development stack", "npm run docker:dev", "Starts database/cache plus watched API/Vite development services through Compose overlays."],
    ["Run migrations", "npm run db:migrate", "Applies committed Drizzle migrations."],
    ["Create demo data", "npm run seed", "Idempotently creates the local household and demo student profile."],
    ["Quality gate", "npm run check", "Runs lint, formatting, type checks, coverage, and web build."],
    ["Production-like stack", "npm run docker:up", "Builds Compose services: PostgreSQL, Redis, one-shot migration, API, and Nginx gateway."],
    ["Public gateway", "Nginx on WEB_PORT (default 8080)", "Serves static web build and proxies /api to the internal Fastify API."],
], [1.45, 2.3, 2.75])

page_break()
heading("8. Current constraints and implementation notes")
heading("Deliberate product and safety limits", 2)
bullets([
    "Deployment mode is closed_demo. Parent enrollment is invitation-only and terms acknowledgement is not verified parental consent.",
    "Deterministic safety matching is intentionally narrow. It is not a substitute for professional safeguarding and will not identify every concerning situation.",
    "Normal evening AI response is fail-closed: absent/unusable AI or encryption configuration returns AI_UNAVAILABLE rather than a generic personalized response.",
    "The parent dashboard is not a transcript dump or report card. Its learning signals are heuristic metadata and must not be presented as mastery or diagnosis.",
    "Transcript access is a separate household-authorized setting; encryption protects stored content but server code with the key can decrypt it, so key storage and rotation remain operational responsibilities.",
])
heading("Engineering follow-through for production", 2)
table(["Area", "Required follow-through"], [
    ["Identity and consent", "Verified parent identity/consent workflow, real account email delivery, account recovery, and jurisdiction-specific legal/privacy review."],
    ["Data lifecycle", "Production-tested export/withdraw/delete behavior, retention policy enforcement, backup/restore testing, and encryption-key rotation plan."],
    ["Network and platform", "TLS-capable edge, explicit CORS policy, trusted-proxy configuration only when correct, secrets management, observability, uptime monitoring, and incident response."],
    ["AI and safety", "Human review of prompt/model/safety-regex changes, adversarial child-safety evaluation, provider data-use approval, and ongoing quality monitoring."],
], [1.65, 4.85])
callout("Use this document", "This reference describes the current repository implementation and its intended product behavior. When making changes, update the relevant feature module, service/schema/test coverage, and the source documentation together so that the parent and child safety boundaries stay explicit.")

# header / footer
for section in doc.sections:
    h = section.header.paragraphs[0]; h.text = "Undercurrent | Complete Application Reference"; h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in h.runs: r.font.name="Aptos"; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
    f = section.footer.paragraphs[0]; f.text = "Internal product and engineering reference"; f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in f.runs: r.font.name="Aptos"; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)

doc.core_properties.title = "Undercurrent Complete Application Reference"
doc.core_properties.subject = "Product, architecture, parent and child flows, and pages"
doc.core_properties.author = "Undercurrent project"
doc.save(OUT)
print(OUT)
